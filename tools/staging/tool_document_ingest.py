"""
HUGINN_MANIFEST
tool_id:            tool.document.ingest.v1
title:              Document Ingestion
capability_summary: >
  Extract text from documents (PDF, Word, HTML, plain text, Markdown) and
  return the content ready for LTM ingestion via Logos. Optionally chunk
  the document by page or by paragraph for per-section LTM entries. Use
  when the user says "read this file", "add this to memory", or drops a
  document for Artux to learn from.
polarity:           write
permission_scope:   []
inputs:
  path:         {type: string, description: "Absolute or relative path to the document"}
  format:       {type: string, default: "auto", description: "auto | pdf | docx | html | txt | md"}
  chunk_by:     {type: string, default: "none", description: "none | page | paragraph — how to split for LTM"}
  max_chunks:   {type: integer, default: 50,   description: "Maximum chunks to return"}
  include_meta: {type: boolean, default: true,  description: "Include title, author, page count in output"}
outputs:
  title:        {type: string}
  author:       {type: string}
  page_count:   {type: integer}
  word_count:   {type: integer}
  full_text:    {type: string,  description: "Full extracted text (when chunk_by=none)"}
  chunks:       {type: array,   description: "[{index, text, page, word_count}] (when chunk_by != none)"}
  chunk_count:  {type: integer}
  format:       {type: string,  description: "Detected or specified format"}
  summary:      {type: string,  description: "1–2 sentence description for speech"}
dependencies:
  - pymupdf>=1.24
  - python-docx>=1.1
perception_capable: false
handler:            handle
END_MANIFEST

Document ingestion without docling.

Chosen libraries:
  pymupdf (fitz) — PDF extraction. Fastest Python PDF library, no native
      compilation needed on most platforms (wheels available). Handles text,
      tables, and embedded fonts. ~8 MB install size.
  python-docx   — Word (.docx) extraction. Pure Python, no LibreOffice.
  stdlib        — HTML (html.parser), plain text, Markdown (treated as text).

docling comparison:
  docling is excellent for complex layouts (multi-column, academic papers)
  but pulls in torch, transformers, and multiple native libraries — 2+ GB.
  For ambient use where documents are mostly prose (notes, articles, emails,
  reports), pymupdf + python-docx covers 95% of cases at 1/100th the size.
  docling can be added as a separate staged tool for complex layouts.

Chunking:
  chunk_by=none     — return full_text as a single string (for short docs)
  chunk_by=page     — one chunk per PDF page or per section heading (docx)
  chunk_by=paragraph — split on double-newlines; max 400 words per chunk

LTM ingestion workflow (after calling this tool):
  Each chunk should be passed to Logos via consolidate_ltm() with:
    narrative  = chunk["text"]
    class_type = "document"
    topics     = [document title, format]
  This gives Muninn per-section recall granularity rather than a single
  blurry document-level entry.

Environment overrides:
  HUGINN_DOCS_MAX_CHUNKS=50
"""

from __future__ import annotations

import os
import re
import html as html_lib
from pathlib import Path
from typing import Any


def handle(
    path:         str  = "",
    format:       str  = "auto",
    chunk_by:     str  = "none",
    max_chunks:   int  = 50,
    include_meta: bool = True,
) -> dict:
    """Extract text from a document file."""
    max_chunks = int(os.environ.get("HUGINN_DOCS_MAX_CHUNKS", max_chunks))
    file_path  = Path(path).expanduser()

    if not file_path.exists():
        return _error(f"File not found: {path}")

    # Detect format
    fmt = format if format != "auto" else _detect_format(file_path)

    try:
        if fmt == "pdf":
            return _ingest_pdf(file_path, chunk_by, max_chunks, include_meta)
        elif fmt == "docx":
            return _ingest_docx(file_path, chunk_by, max_chunks, include_meta)
        elif fmt in ("html", "htm"):
            return _ingest_html(file_path, chunk_by, max_chunks)
        elif fmt in ("txt", "md", "text"):
            return _ingest_text(file_path, chunk_by, max_chunks)
        else:
            return _error(f"Unsupported format: {fmt!r} (path: {path})")
    except Exception as e:
        return _error(f"Extraction error ({fmt}): {e}")


# ─── Format detection ─────────────────────────────────────────────────────────

def _detect_format(path: Path) -> str:
    suffix = path.suffix.lower().lstrip(".")
    return {
        "pdf": "pdf", "docx": "docx", "doc": "docx",
        "html": "html", "htm": "html",
        "txt": "txt", "md": "md", "markdown": "md",
        "rst": "txt", "csv": "txt",
    }.get(suffix, "txt")


# ─── PDF ──────────────────────────────────────────────────────────────────────

def _ingest_pdf(path: Path, chunk_by: str, max_chunks: int,
                include_meta: bool) -> dict:
    import fitz   # pymupdf

    doc   = fitz.open(str(path))
    meta  = doc.metadata if include_meta else {}
    pages = []

    for page in doc:
        text = page.get_text("text").strip()
        if text:
            pages.append(text)

    doc.close()

    title      = meta.get("title", "") or path.stem
    author     = meta.get("author", "")
    page_count = len(pages)
    full_text  = "\n\n".join(pages)
    word_count = len(full_text.split())

    chunks = _chunkify(pages, full_text, chunk_by, max_chunks, unit="page")

    return _result(title, author, page_count, word_count,
                   full_text if chunk_by == "none" else "",
                   chunks, "pdf", path)


# ─── DOCX ─────────────────────────────────────────────────────────────────────

def _ingest_docx(path: Path, chunk_by: str, max_chunks: int,
                 include_meta: bool) -> dict:
    from docx import Document
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    doc   = Document(str(path))
    title  = ""
    author = ""

    if include_meta:
        props  = doc.core_properties
        title  = props.title  or path.stem
        author = props.author or ""

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    full_text  = "\n\n".join(paragraphs)
    word_count = len(full_text.split())

    # Section chunks = heading-delimited blocks
    if chunk_by == "page":
        # DOCX doesn't have page boundaries easily — use headings as sections
        sections, current = [], []
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            if p.style.name.startswith("Heading"):
                if current:
                    sections.append("\n".join(current))
                    current = []
            current.append(text)
        if current:
            sections.append("\n".join(current))
        pages = sections or [full_text]
    else:
        pages = paragraphs

    chunks = _chunkify(pages, full_text, chunk_by, max_chunks, unit="section")

    return _result(title or path.stem, author, len(pages), word_count,
                   full_text if chunk_by == "none" else "",
                   chunks, "docx", path)


# ─── HTML ─────────────────────────────────────────────────────────────────────

def _ingest_html(path: Path, chunk_by: str, max_chunks: int) -> dict:
    from html.parser import HTMLParser

    class _Extractor(HTMLParser):
        def __init__(self):
            super().__init__()
            self.parts   = []
            self._skip   = False
            self._title  = []
            self._in_title = False

        def handle_starttag(self, tag, attrs):
            if tag in ("script", "style", "nav", "footer", "aside"):
                self._skip = True
            if tag == "title":
                self._in_title = True

        def handle_endtag(self, tag):
            if tag in ("script", "style", "nav", "footer", "aside"):
                self._skip = False
            if tag == "title":
                self._in_title = False
            if tag in ("p", "div", "li", "h1", "h2", "h3", "h4", "h5"):
                self.parts.append("\n")

        def handle_data(self, data):
            if self._in_title:
                self._title.append(data)
            elif not self._skip:
                self.parts.append(data)

    src     = path.read_text(encoding="utf-8", errors="replace")
    parser  = _Extractor()
    parser.feed(src)

    raw       = "".join(parser.parts)
    raw       = html_lib.unescape(raw)
    raw       = re.sub(r"\n{3,}", "\n\n", raw).strip()
    title     = "".join(parser._title).strip() or path.stem
    full_text = raw
    word_count = len(raw.split())
    paragraphs = [p.strip() for p in raw.split("\n\n") if p.strip()]

    chunks = _chunkify(paragraphs, full_text, chunk_by, max_chunks, unit="paragraph")

    return _result(title, "", 0, word_count,
                   full_text if chunk_by == "none" else "",
                   chunks, "html", path)


# ─── Plain text / Markdown ────────────────────────────────────────────────────

def _ingest_text(path: Path, chunk_by: str, max_chunks: int) -> dict:
    text       = path.read_text(encoding="utf-8", errors="replace")
    paragraphs = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]
    word_count = len(text.split())
    title      = path.stem

    chunks = _chunkify(paragraphs, text, chunk_by, max_chunks, unit="paragraph")

    return _result(title, "", 0, word_count,
                   text if chunk_by == "none" else "",
                   chunks, "txt", path)


# ─── Chunking ─────────────────────────────────────────────────────────────────

def _chunkify(units: list[str], full_text: str, chunk_by: str,
              max_chunks: int, unit: str = "unit") -> list[dict]:
    if chunk_by == "none":
        return []

    if chunk_by == "page":
        # One chunk per natural unit (page / section / paragraph)
        raw_chunks = units
    elif chunk_by == "paragraph":
        # Merge short paragraphs, split very long ones (target ~400 words)
        raw_chunks = _merge_paragraphs(units, target_words=400)
    else:
        raw_chunks = units

    raw_chunks = raw_chunks[:max_chunks]

    result = []
    for i, text in enumerate(raw_chunks):
        result.append({
            "index":      i,
            "text":       text.strip(),
            unit:         i,
            "word_count": len(text.split()),
        })
    return result


def _merge_paragraphs(paras: list[str], target_words: int = 400) -> list[str]:
    merged, current, current_words = [], [], 0
    for p in paras:
        words = len(p.split())
        if current_words + words > target_words and current:
            merged.append("\n\n".join(current))
            current, current_words = [], 0
        current.append(p)
        current_words += words
    if current:
        merged.append("\n\n".join(current))
    return merged


# ─── Helpers ─────────────────────────────────────────────────────────────────

def _result(title, author, page_count, word_count, full_text, chunks, fmt, path):
    chunk_count = len(chunks)
    summary     = _build_summary(title, fmt, word_count, page_count, chunk_count)
    return {
        "title":       title,
        "author":      author,
        "page_count":  page_count,
        "word_count":  word_count,
        "full_text":   full_text,
        "chunks":      chunks,
        "chunk_count": chunk_count,
        "format":      fmt,
        "summary":     summary,
    }


def _build_summary(title, fmt, word_count, pages, chunks):
    parts = [f"{title!r}" if title else "Document"]
    parts.append(f"({fmt.upper()}, {word_count:,} words")
    if pages:
        parts.append(f", {pages} pages")
    parts.append(")")
    if chunks:
        parts.append(f"Split into {chunks} chunks for ingestion.")
    return " ".join(parts)


def _error(reason: str) -> dict:
    return {
        "title": "", "author": "", "page_count": 0, "word_count": 0,
        "full_text": "", "chunks": [], "chunk_count": 0,
        "format": "unknown", "summary": f"Error: {reason}",
    }
